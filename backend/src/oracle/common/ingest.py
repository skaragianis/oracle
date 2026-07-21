import logging
import mimetypes
import os
import re
import shutil
import sqlite3
import uuid
from collections.abc import Callable, Iterable, Iterator
from dataclasses import dataclass
from pathlib import Path
from typing import NamedTuple

import docx
import fitz
import tiktoken
from docx.table import Table
from docx.text.paragraph import Paragraph

from oracle.common import chunks, documents, embeddings

logger = logging.getLogger(__name__)

SUPPORTED_SUFFIXES = {".pdf", ".docx"}

DEFAULT_UPLOADS_DIR = Path(
    os.environ.get("ORACLE_UPLOADS_DIR")
    or Path(__file__).resolve().parents[3] / "data" / "uploads"
)

CHUNK_ENCODING_NAME = "o200k_base"
CHUNK_TARGET_TOKENS = 800
CHUNK_OVERLAP_TOKENS = 100


class UnsupportedFileTypeError(ValueError):
    pass


ShouldStop = Callable[[], bool]


class IngestInterrupted(Exception):
    pass


def _check_stop(should_stop: ShouldStop | None) -> None:
    if should_stop is not None and should_stop():
        raise IngestInterrupted


@dataclass
class StagedDocument:
    doc_id: int
    destination_path: Path
    replaced: bool


@dataclass
class ProcessResult:
    status: str
    error: str | None


@dataclass
class IngestResult:
    doc_id: int
    destination_path: Path
    replaced: bool
    status: str
    error: str | None


def stage_file(
    conn: sqlite3.Connection,
    source_path: str | Path,
    uploads_dir: Path = DEFAULT_UPLOADS_DIR,
    vector_index: embeddings.VectorIndex | None = None,
) -> StagedDocument:
    source_path = Path(source_path)

    if source_path.suffix.lower() not in SUPPORTED_SUFFIXES:
        raise UnsupportedFileTypeError(
            f"Unsupported file type {source_path.suffix!r}; "
            f"expected one of {sorted(SUPPORTED_SUFFIXES)}"
        )

    if not source_path.is_file():
        raise FileNotFoundError(source_path)

    existing = documents.find_document_by_filename(conn, source_path.name)

    uploads_dir.mkdir(parents=True, exist_ok=True)
    destination_path = uploads_dir / f"{uuid.uuid4()}{source_path.suffix.lower()}"
    shutil.copy2(source_path, destination_path)

    mime_type, _ = mimetypes.guess_type(source_path.name)

    if existing is not None:
        doc_id, previous_stored_filename = existing
        chunks.delete_chunks_for_document(conn, doc_id)
        if vector_index is not None:
            vector_index.delete_document(doc_id)
        documents.replace_document_upload(
            conn,
            doc_id,
            stored_filename=destination_path.name,
            mime_type=mime_type or "application/octet-stream",
            size_bytes=destination_path.stat().st_size,
        )
        (uploads_dir / previous_stored_filename).unlink(missing_ok=True)
    else:
        doc_id = documents.create_document(
            conn,
            filename=source_path.name,
            stored_filename=destination_path.name,
            mime_type=mime_type or "application/octet-stream",
            size_bytes=destination_path.stat().st_size,
        )

    return StagedDocument(
        doc_id=doc_id,
        destination_path=destination_path,
        replaced=existing is not None,
    )


def process_document(
    conn: sqlite3.Connection,
    doc_id: int,
    stored_path: str | Path,
    vector_index: embeddings.VectorIndex | None = None,
    should_stop: ShouldStop | None = None,
) -> ProcessResult:
    stored_path = Path(stored_path)

    try:
        # May occur if previously interrupted
        chunks.delete_chunks_for_document(conn, doc_id)
        if vector_index is not None:
            vector_index.delete_document(doc_id)
        chunk_document(conn, doc_id, stored_path, should_stop=should_stop)
        if vector_index is not None:
            _embed_document_chunks(conn, doc_id, vector_index, should_stop=should_stop)
    except IngestInterrupted:
        logger.info("Ingest of document %s interrupted; left pending", doc_id)
        return ProcessResult(status="pending", error=None)
    except Exception as exc:
        logger.exception("Failed to process document %s at %s", doc_id, stored_path)
        error = f"{type(exc).__name__}: {exc}"
        documents.mark_document_failed(conn, doc_id, error)
        return ProcessResult(status="failed", error=error)

    documents.mark_document_ready(conn, doc_id)
    return ProcessResult(status="ready", error=None)


def _embed_document_chunks(
    conn: sqlite3.Connection,
    doc_id: int,
    vector_index: embeddings.VectorIndex,
    should_stop: ShouldStop | None = None,
) -> None:
    rows = conn.execute(
        "SELECT id, text FROM chunks WHERE doc_id = ? ORDER BY seq", (doc_id,)
    ).fetchall()
    for start in range(0, len(rows), embeddings.EMBED_BATCH_SIZE):
        _check_stop(should_stop)
        batch = rows[start : start + embeddings.EMBED_BATCH_SIZE]
        vector_index.index_chunks(
            doc_id,
            [embeddings.ChunkToIndex(chunk_id=row[0], text=row[1]) for row in batch],
        )


def ingest_file(
    conn: sqlite3.Connection,
    source_path: str | Path,
    uploads_dir: Path = DEFAULT_UPLOADS_DIR,
    vector_index: embeddings.VectorIndex | None = None,
) -> IngestResult:
    staged = stage_file(
        conn, source_path, uploads_dir=uploads_dir, vector_index=vector_index
    )
    processed = process_document(
        conn, staged.doc_id, staged.destination_path, vector_index=vector_index
    )

    return IngestResult(
        doc_id=staged.doc_id,
        destination_path=staged.destination_path,
        replaced=staged.replaced,
        status=processed.status,
        error=processed.error,
    )


@dataclass
class _BufferedLine:
    text: str
    page_number: int | None
    paragraph_index: int
    char_start: int
    token_count: int


class LineSegment(NamedTuple):
    text: str
    page_number: int | None
    paragraph_index: int


def chunk_document(
    conn: sqlite3.Connection,
    doc_id: int,
    path: str | Path,
    *,
    encoding_name: str = CHUNK_ENCODING_NAME,
    target_tokens: int = CHUNK_TARGET_TOKENS,
    overlap_tokens: int = CHUNK_OVERLAP_TOKENS,
    should_stop: ShouldStop | None = None,
) -> int:
    path = Path(path)
    suffix = path.suffix.lower()
    lines: Iterator[LineSegment]
    if suffix == ".pdf":
        lines = _iter_pdf_lines(path)
    elif suffix == ".docx":
        lines = _iter_docx_lines(path)
    else:
        raise UnsupportedFileTypeError(f"Cannot chunk {suffix} documents yet")

    return _chunk_lines(
        conn,
        doc_id,
        lines,
        encoding_name=encoding_name,
        target_tokens=target_tokens,
        overlap_tokens=overlap_tokens,
        should_stop=should_stop,
    )


_HYPHENATED_BREAK_RE = re.compile(r"(\w)-\n(\w)")
_WHITESPACE_RUN_RE = re.compile(r"[ \t]+")
_BLANK_LINE_RE = re.compile(r"\n[ \t]*\n+")


def _split_into_paragraphs(block_text: str) -> list[str]:
    text = _HYPHENATED_BREAK_RE.sub(r"\1\2", block_text)
    paragraphs = []
    for raw_paragraph in _BLANK_LINE_RE.split(text):
        collapsed = raw_paragraph.replace("\n", " ")
        normalized = _WHITESPACE_RUN_RE.sub(" ", collapsed).strip()
        if normalized:
            paragraphs.append(normalized)
    return paragraphs


def _iter_pdf_lines(pdf_path: Path) -> Iterator[LineSegment]:
    paragraph_index = -1
    with fitz.open(pdf_path) as doc:
        for page_number, page in enumerate(doc, start=1):
            for block in page.get_text("blocks"):
                block_text = block[4]
                if not block_text.strip():
                    continue
                for paragraph in _split_into_paragraphs(block_text):
                    paragraph_index += 1
                    yield LineSegment(paragraph + "\n", page_number, paragraph_index)


def _iter_docx_lines(docx_path: Path) -> Iterator[LineSegment]:
    paragraph_index = -1
    for block in docx.Document(str(docx_path)).iter_inner_content():
        for paragraph in _iter_block_paragraphs(block):
            text = paragraph.text.strip()
            if not text:
                continue
            paragraph_index += 1
            yield LineSegment(text + "\n", None, paragraph_index)


def _iter_block_paragraphs(block: Paragraph | Table) -> Iterator[Paragraph]:
    if isinstance(block, Table):
        for row in block.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    else:
        yield block


def _chunk_lines(
    conn: sqlite3.Connection,
    doc_id: int,
    lines: Iterable[LineSegment],
    *,
    encoding_name: str,
    target_tokens: int,
    overlap_tokens: int,
    should_stop: ShouldStop | None = None,
) -> int:
    encoding = tiktoken.get_encoding(encoding_name)
    seq = 0
    char_offset = 0
    buffer: list[_BufferedLine] = []
    buffer_tokens = 0
    buffer_has_new_content = False

    def flush() -> None:
        nonlocal seq, buffer, buffer_tokens, buffer_has_new_content
        if not buffer_has_new_content:
            return

        first, last = buffer[0], buffer[-1]
        chunks.create_chunk(
            conn,
            doc_id=doc_id,
            seq=seq,
            text="".join(line.text for line in buffer),
            page_number=first.page_number,
            paragraph_start=first.paragraph_index,
            paragraph_end=last.paragraph_index,
            char_start=first.char_start,
            char_end=last.char_start + len(last.text),
        )
        seq += 1

        overlap: list[_BufferedLine] = []
        overlap_tokens_used = 0
        for line in reversed(buffer):
            if overlap and overlap_tokens_used + line.token_count > overlap_tokens:
                break
            overlap.insert(0, line)
            overlap_tokens_used += line.token_count
            if overlap_tokens_used >= overlap_tokens:
                break

        buffer = overlap
        buffer_tokens = overlap_tokens_used
        buffer_has_new_content = False

    for line_text, page_number, paragraph_index in lines:
        _check_stop(should_stop)
        token_count = len(encoding.encode(line_text))

        buffer.append(
            _BufferedLine(
                text=line_text,
                page_number=page_number,
                paragraph_index=paragraph_index,
                char_start=char_offset,
                token_count=token_count,
            )
        )
        buffer_tokens += token_count
        buffer_has_new_content = True
        char_offset += len(line_text)

        if buffer_tokens >= target_tokens:
            flush()

    flush()
    return seq


def delete_document(
    conn: sqlite3.Connection,
    doc_id: int,
    uploads_dir: Path = DEFAULT_UPLOADS_DIR,
    vector_index: embeddings.VectorIndex | None = None,
) -> bool:
    stored_filename = documents.get_stored_filename(conn, doc_id)
    if stored_filename is None:
        return False
    chunks.delete_chunks_for_document(conn, doc_id)
    if vector_index is not None:
        vector_index.delete_document(doc_id)
    documents.delete_document(conn, doc_id)
    (uploads_dir / stored_filename).unlink(missing_ok=True)
    return True


def remove_document_uploads() -> None:
    for item in DEFAULT_UPLOADS_DIR.iterdir():
        item.unlink()
